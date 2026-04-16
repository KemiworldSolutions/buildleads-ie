"""
Format-aware text readers for PDF, DOCX, and legacy DOC council files.

DOC (legacy Word binary) is the ugly one. Pure Python .doc parsing is
non-trivial, so we use a pragmatic approach: open the WordDocument stream
with olefile and extract printable UTF-16/ASCII runs. The output is noisier
than a clean parse, but it's good enough for an LLM to extract structured
data from — the LLM filters the garbage.
"""

from __future__ import annotations

import re
from pathlib import Path

from pypdf import PdfReader


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def read_docx(path: Path) -> str:
    from docx import Document  # python-docx
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def read_doc(path: Path) -> str:
    """Extract text from a legacy MS Word .doc (OLE) file."""
    import olefile
    ole = olefile.OleFileIO(str(path))
    if not ole.exists("WordDocument"):
        return ""
    stream = ole.openstream("WordDocument").read()
    # Try UTF-16-LE first (Word stores text as UTF-16 by default)
    text_u16 = stream.decode("utf-16-le", errors="replace")
    runs_u16 = re.findall(r"[\x20-\x7e\n\r\t]{8,}", text_u16)
    # Also try raw ASCII in case the document has any CP1252 runs
    text_ascii = stream.decode("latin-1", errors="replace")
    runs_ascii = re.findall(r"[\x20-\x7e\n\r\t]{8,}", text_ascii)
    # Take whichever gave more text
    u16 = "\n".join(runs_u16)
    a = "\n".join(runs_ascii)
    chosen = u16 if len(u16) > len(a) else a
    # Collapse excessive whitespace runs
    chosen = re.sub(r"[ \t]{3,}", "  ", chosen)
    chosen = re.sub(r"\n{3,}", "\n\n", chosen)
    return chosen


def read_any(path: Path, fmt: str) -> str:
    fmt = fmt.lower()
    if fmt == "pdf":
        return read_pdf(path)
    if fmt == "docx":
        return read_docx(path)
    if fmt == "doc":
        return read_doc(path)
    raise ValueError(f"Unknown format: {fmt}")
